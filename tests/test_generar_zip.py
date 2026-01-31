"""Tests for generar_zip() — legal document and ZIP generation."""

import zipfile
from io import BytesIO

from docx import Document

from app import generar_zip


class TestGenerarZipStructure:
    """Verify ZIP file structure and contents."""

    def test_returns_readable_zip(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        assert isinstance(result, BytesIO)
        with zipfile.ZipFile(result, "r") as zf:
            assert zf.testzip() is None  # no corrupt files

    def test_contains_instructions_file(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        with zipfile.ZipFile(result, "r") as zf:
            names = zf.namelist()
            assert "INSTRUCCIONES.txt" in names

    def test_instructions_content(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        with zipfile.ZipFile(result, "r") as zf:
            content = zf.read("INSTRUCCIONES.txt").decode("utf-8")
            assert "Imprime" in content
            assert "Firma" in content

    def test_one_docx_per_court(self, sample_datos, sample_multas):
        """Two fines at the same court -> one .docx file."""
        result = generar_zip(sample_datos, sample_multas)
        with zipfile.ZipFile(result, "r") as zf:
            docx_files = [n for n in zf.namelist() if n.endswith(".docx")]
            assert len(docx_files) == 1

    def test_multiple_courts_produce_multiple_docx(
        self, sample_datos, sample_multas_multiple_courts
    ):
        """Fines at two different courts -> two .docx files."""
        result = generar_zip(sample_datos, sample_multas_multiple_courts)
        with zipfile.ZipFile(result, "r") as zf:
            docx_files = [n for n in zf.namelist() if n.endswith(".docx")]
            assert len(docx_files) == 2

    def test_docx_filename_contains_plate(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        with zipfile.ZipFile(result, "r") as zf:
            docx_files = [n for n in zf.namelist() if n.endswith(".docx")]
            assert any(sample_datos["patente"] in name for name in docx_files)

    def test_docx_filename_starts_with_escrito_jpl(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        with zipfile.ZipFile(result, "r") as zf:
            docx_files = [n for n in zf.namelist() if n.endswith(".docx")]
            assert all(name.startswith("Escrito JPL") for name in docx_files)


class TestGenerarZipDocumentContent:
    """Verify the generated Word document has correct legal content."""

    def _read_docx_from_zip(self, zip_buffer):
        """Extract the first .docx from the ZIP and return a Document object."""
        with zipfile.ZipFile(zip_buffer, "r") as zf:
            docx_name = [n for n in zf.namelist() if n.endswith(".docx")][0]
            docx_bytes = zf.read(docx_name)
        return Document(BytesIO(docx_bytes))

    def test_document_contains_owner_name(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert sample_datos["nombre"] in full_text

    def test_document_contains_rut(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert sample_datos["rut"] in full_text

    def test_document_contains_patente(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert sample_datos["patente"] in full_text

    def test_document_contains_legal_article_reference(
        self, sample_datos, sample_multas
    ):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "artículo 24" in full_text or "Art. 24" in full_text

    def test_document_contains_law_reference(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "18.287" in full_text

    def test_document_has_table_with_fines(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        # Header row + 2 fine rows
        assert len(table.rows) == 3

    def test_table_header_labels(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        table = doc.tables[0]
        headers = [cell.text for cell in table.rows[0].cells]
        assert "ROL CAUSA" in headers
        assert "FECHA INGRESO" in headers

    def test_table_contains_fine_roles(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        table = doc.tables[0]
        all_cell_text = []
        for row in table.rows[1:]:
            for cell in row.cells:
                all_cell_text.append(cell.text)
        assert "123-2020" in all_cell_text
        assert "456-2021" in all_cell_text

    def test_document_has_signature_section(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "FIRMA PROPIETARIO" in full_text

    def test_document_has_otrosies(self, sample_datos, sample_multas):
        result = generar_zip(sample_datos, sample_multas)
        doc = self._read_docx_from_zip(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "PRIMER OTROSÍ" in full_text
        assert "SEGUNDO OTROSÍ" in full_text
