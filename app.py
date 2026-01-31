import streamlit as st
import pdfplumber
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import zipfile

# --- CONFIGURACIÓN DE NEGOCIO ---
PRECIO_SERVICIO = 14990
# ⚠️ IMPORTANTE: Pega aquí tu link real de Mercado Pago
LINK_MERCADO_PAGO = "https://mpago.la/24oHCqt" 
CLAVE_ACCESO = "AUTO2026"

# --- CONFIGURACIÓN VISUAL ---
_PAGE_STYLE = """
    <style>
    #MainMenu {visibility: hidden;} footer {visibility: hidden;} header {visibility: hidden;}
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;900&display=swap');
    html, body, [class*="css"]  {font-family: 'Inter', sans-serif; background-color: #f8fafc;}

    .hero {text-align: center; padding: 30px 0;}
    .hero h1 {color: #0f172a; font-weight: 900; font-size: 2.5rem; letter-spacing: -1px; margin-bottom: 5px;}

    .instruction-box {
        background: white; padding: 25px; border-radius: 12px;
        border: 1px solid #e2e8f0; text-align: center; margin-bottom: 20px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
        height: 100%;
    }
    .btn-registrocivil {
        display: block; width: 100%; padding: 15px; margin-top: 15px;
        background-color: #00519E; color: white; text-decoration: none; border-radius: 8px; text-align: center; font-weight: bold; border: 1px solid #003f7a; transition: 0.3s;
    }
    .pay-btn {
        display: block; width: 100%; background: #FACC15; color: black;
        font-weight: 900; text-align: center; padding: 20px; border-radius: 12px;
        text-decoration: none; font-size: 1.4rem; box-shadow: 0 10px 15px -3px rgba(250, 204, 21, 0.3);
        border: 2px solid #eab308;
    }
    .success-box {
        background: white; border: 2px solid #22c55e; border-radius: 16px;
        padding: 30px; text-align: center; margin-top: 20px;
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
    }
    .money-tag {
        font-size: 3.5rem; font-weight: 900; color: #15803d;
        letter-spacing: -2px; line-height: 1.1; margin: 15px 0;
    }
    </style>
    """

# --- LÓGICA DE EXTRACCIÓN ---

def limpiar_texto(texto):
    if not texto: return ""
    # Eliminamos comillas y espacios extra, mantenemos mayúsculas
    return texto.replace('"', '').replace(',', '').strip().upper()

def limpiar_juzgado(nombre_juzgado):
    """Limpia el nombre del juzgado para el nombre del archivo."""
    nombre = nombre_juzgado.upper()
    nombre = nombre.replace("JUZGADO DE POLICIA LOCAL", "").replace("JUZGADO POLICIA LOCAL", "")
    nombre = nombre.replace("TRIBUNAL", "").strip()
    return nombre

def buscar_patente_universal(texto):
    # 1. Búsqueda por etiqueta explícita (PLACA PATENTE...)
    match_etiqueta = re.search(r'(?:PLACA|PATENTE|PPU).*?([A-Z0-9]{2,4}[\s\.-]?\d{2,4})', texto)
    if match_etiqueta:
        raw = match_etiqueta.group(1).replace(".", "").replace(" ", "").replace("-", "")
        if len(raw) == 6: return raw

    # 2. Búsqueda libre en cabecera (Formato Nuevo y Antiguo)
    cabecera = texto[:1000]
    
    patron_nueva = re.search(r'\b([B-D,F-H,J-L,P,R-T,V-Z]{4})[\s\.-]?(\d{2})\b', cabecera)
    if patron_nueva: return f"{patron_nueva.group(1)}{patron_nueva.group(2)}"
    
    patron_antigua = re.search(r'\b([A-Z]{2})[\s\.-]?(\d{4})\b', cabecera)
    if patron_antigua: return f"{patron_antigua.group(1)}{patron_antigua.group(2)}"
        
    return "NO_DETECTADA"

def es_prescribible(fecha_str):
    try:
        fecha_clean = fecha_str.split(" ")[0].strip()
        fecha_obj = datetime.strptime(fecha_clean, '%d-%m-%Y')
        hoy = datetime.now()
        return (hoy - fecha_obj).days > 1095
    except:
        return False

def procesar_pdf(archivo):
    multas = []
    # Valores por defecto para que no falle si no encuentra algo
    datos = {"patente": "NO_DETECTADA", "rut": "NO DETECTADO", "nombre": "PROPIETARIO"}
    try:
        with pdfplumber.open(archivo) as pdf:
            texto_completo = ""
            for page in pdf.pages: texto_completo += page.extract_text() + "\n"
            
        texto_limpio = texto_completo.replace("\n", " ")

        # Validación básica de documento
        if "REGISTRO DE MULTAS" not in texto_limpio and "TRANSITO NO PAGADAS" not in texto_limpio:
            return None, None

        # --- EXTRACCIÓN DE DATOS ---
        datos['patente'] = buscar_patente_universal(texto_limpio)
        
        match_rut = re.search(r'R\.U\.N\.\s*:\s*([\d\.\-Kk]+)', texto_completo)
        if match_rut: datos['rut'] = limpiar_texto(match_rut.group(1))
        
        # --- CORRECCIÓN NOMBRE (Nueva estrategia) ---
        # 1. Intento principal: Buscar la línea que empieza con "Nombre :"
        match_nombre = re.search(r'(?:Nombre|Propietario)\s*[:\.]\s*([A-ZÑ\s]+)', texto_completo, re.IGNORECASE)
        if match_nombre:
            datos['nombre'] = limpiar_texto(match_nombre.group(1))
        else:
            # 2. Intento secundario (si el formato es distinto)
            match_nombre_alt = re.search(r'\n([A-ZÑ\s]{10,})\n', texto_completo) # Busca una línea con solo texto en mayúsculas (nombre probable)
            if match_nombre_alt:
                # Verificamos que no sea un título
                posible_nombre = match_nombre_alt.group(1)
                if "REGISTRO" not in posible_nombre and "CERTIFICADO" not in posible_nombre:
                    datos['nombre'] = limpiar_texto(posible_nombre)

        # --- EXTRACCIÓN DE MULTAS ---
        bloques = texto_completo.split("ID MULTA")
        for bloque in bloques:
            if "TRIBUNAL" in bloque:
                juzgado_match = re.search(r'TRIBUNAL\s*:\s*(.+)', bloque)
                rol_match = re.search(r'(?<!AÑO )ROL\s*:\s*([\w\-\.]+)', bloque)
                fecha_ingreso_match = re.search(r'FECHA INGRESO RMNP\s*:\s*([\d\-\s:]+)', bloque)
                
                if juzgado_match and rol_match and fecha_ingreso_match:
                    fecha_ingreso = fecha_ingreso_match.group(1).strip()
                    if es_prescribible(fecha_ingreso):
                        multas.append({
                            "juzgado": juzgado_match.group(1).strip(),
                            "rol": rol_match.group(1).strip(),
                            "fecha_ingreso": fecha_ingreso.split(" ")[0]
                        })
        return datos, multas
    except Exception as e:
        return None, None

# --- GENERADOR WORD (FORMATO LEGAL JUDICIAL) ---
def generar_zip(datos, multas):
    multas_por_juzgado = {}
    for m in multas:
        jz = m['juzgado']
        if jz not in multas_por_juzgado: multas_por_juzgado[jz] = []
        multas_por_juzgado[jz].append(m)
    
    memoria_zip = BytesIO()
    with zipfile.ZipFile(memoria_zip, 'w') as zf:
        for juzgado, lista_multas in multas_por_juzgado.items():
            doc = Document()
            
            # Configurar márgenes estándar (Carta)
            section = doc.sections[0]
            section.page_height = Mm(279.4)
            section.page_width = Mm(215.9)
            section.left_margin = Mm(25.4)
            section.right_margin = Mm(25.4)
            section.top_margin = Mm(25.4)
            section.bottom_margin = Mm(25.4)
            
            # Fuente Arial 12
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)
            
            # 1. LA SUMA (Alineada Derecha)
            p_suma = doc.add_paragraph()
            p_suma.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_suma.add_run("EN LO PRINCIPAL: ").bold = True
            p_suma.add_run("Solicita Prescripción de Multas (Art. 24 Ley 18.287).\n")
            p_suma.add_run("PRIMER OTROSÍ: ").bold = True
            p_suma.add_run("Acompaña documento.\n")
            p_suma.add_run("SEGUNDO OTROSÍ: ").bold = True
            p_suma.add_run("Solicita notificación por correo electrónico.")

            doc.add_paragraph() 

            # 2. EL ENCABEZADO (Centrado y Negrita)
            p_sjl = doc.add_paragraph()
            p_sjl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sjl = p_sjl.add_run(f"S.J.L. DE {juzgado}")
            run_sjl.bold = True
            
            doc.add_paragraph()

            # 3. CUERPO PRINCIPAL (Justificado)
            p_cuerpo = doc.add_paragraph()
            p_cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- AQUÍ ESTÁ EL CAMBIO IMPORTANTE
            
            p_cuerpo.add_run(f"{datos['nombre']}").bold = True
            p_cuerpo.add_run(f", cédula nacional de identidad N° {datos['rut']}, domiciliado en ")
            p_cuerpo.add_run("__________________________________________________________").bold = True 
            p_cuerpo.add_run(", comuna de _______________, en los autos sobre infracción a la Ley de Tránsito, placa patente única ")
            p_cuerpo.add_run(f"{datos['patente']}").bold = True
            p_cuerpo.add_run(", a US. respetuosamente digo:")
            
            doc.add_paragraph()

            p_argumento = doc.add_paragraph("Que, por este acto, vengo en solicitar se declare la prescripción de las multas empadronadas que se detallan a continuación, en razón de lo dispuesto en el artículo 24 de la Ley N° 18.287, por haber transcurrido más de tres años desde su anotación en el Registro de Multas de Tránsito no Pagadas:")
            p_argumento.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justificado

            doc.add_paragraph()

            # 4. TABLA DE MULTAS (Centrada)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ROL CAUSA'
            hdr_cells[1].text = 'FECHA INGRESO'
            
            # Negrita en encabezados
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True

            # Llenar datos
            for m in lista_multas:
                row_cells = table.add_row().cells
                row_cells[0].text = str(m['rol'])
                row_cells[1].text = str(m['fecha_ingreso'])
                # Centrar contenido
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # 5. PETITORIO (Justificado)
            p_tanto = doc.add_paragraph()
            p_tanto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_tanto.add_run("POR TANTO, ").bold = True
            p_tanto.add_run("con el mérito de lo expuesto y del tiempo transcurrido,")
            
            p_ruego = doc.add_paragraph("RUEGO A US. acceder a lo solicitado, declarando la prescripción de la(s) multa(s) individualizada(s) y ordenando su eliminación del Registro.")
            p_ruego.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_paragraph()
            
            # 6. OTROSÍES (Justificado)
            p_otrosi1 = doc.add_paragraph()
            p_otrosi1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_otrosi1.add_run("PRIMER OTROSÍ: ").bold = True
            p_otrosi1.add_run("Sírvase US. tener por acompañado el Certificado de Multas de Tránsito no Pagadas emitido por el Servicio de Registro Civil e Identificación.")
            
            p_otrosi2 = doc.add_paragraph()
            p_otrosi2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_otrosi2.add_run("SEGUNDO OTROSÍ: ").bold = True
            p_otrosi2.add_run("Vengo en solicitar se me notifique la resolución de esta solicitud al correo electrónico: ")
            p_otrosi2.add_run("____________________________________________________").bold = True
            
            # 7. FIRMA (Centrada)
            doc.add_paragraph("\n\n\n___________________________\nFIRMA PROPIETARIO").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"{datos['nombre']}\nR.U.N: {datos['rut']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # GUARDAR
            juzgado_limpio = limpiar_juzgado(juzgado)
            nombre_archivo = f"Escrito JPL {juzgado_limpio}_{datos['patente']}.docx"
            
            doc_io = BytesIO(); doc.save(doc_io); doc_io.seek(0)
            zf.writestr(nombre_archivo, doc_io.getvalue())
        
        zf.writestr("INSTRUCCIONES.txt", "1. Imprime 3 copias.\n2. Rellena domicilio y correo.\n3. Firma.\n4. Adjunta el Certificado.")
    memoria_zip.seek(0)
    return memoria_zip

# --- FRONTEND WEB ---

def main():
    st.set_page_config(page_title="BorraTusMultas.cl", page_icon="⚖️", layout="centered")
    st.markdown(_PAGE_STYLE, unsafe_allow_html=True)
    st.markdown("""
    <div class="hero">
        <h1>⚖️ BorraTusMultas.cl</h1>
        <p>Detector Automático de Prescripción</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        st.info("PASO 1: Compra tu Certificado ($1.310) en registrocivil.cl")
        st.markdown(f'<a href="https://www.registrocivil.cl/principal/servicios-en-linea" target="_blank" class="btn-registrocivil">Ir a RegistroCivil.cl</a>', unsafe_allow_html=True)

    with col2:
        st.info("PASO 2: Sube el PDF para analizar")
        uploaded_file = st.file_uploader("Carga el archivo aquí", type="pdf")

    if uploaded_file:
        with st.spinner('Analizando documento...'):
            datos, multas = procesar_pdf(uploaded_file)

        if datos is None:
            st.error("⚠️ Archivo no válido. Sube el Certificado de Multas original.")

        elif multas:
            ahorro = len(multas) * 65000
            st.markdown(f"""
            <div class="success-box">
                <h2>¡{len(multas)} MULTAS BORRABLES DETECTADAS!</h2>
                <p>Vehículo: <b>{datos['patente']}</b></p>
                <div class="money-tag">${ahorro:,.0f}</div>
                <p>AHORRO ESTIMADO</p>
            </div>
            """, unsafe_allow_html=True)

            c1, c2 = st.columns([1.2, 1])
            with c1:
                st.write(" ")
                st.markdown(f'<a href="{LINK_MERCADO_PAGO}" target="_blank" class="pay-btn">DESCARGAR ESCRITOS<br><span style="font-size:1rem; font-weight:normal">${PRECIO_SERVICIO:,.0f}</span></a>', unsafe_allow_html=True)

            with c2:
                st.write(" ")
                clave = st.text_input("Ingresa tu clave de pago:", placeholder="Ej: AUTO2026")
                if clave == CLAVE_ACCESO:
                    zip_buffer = generar_zip(datos, multas)
                    st.balloons()
                    st.download_button("📥 DESCARGAR ZIP", zip_buffer, f"Pack_Legal_{datos['patente']}.zip", "application/zip")
                elif clave:
                    st.error("Clave incorrecta.")

        else:
            st.warning(f"Estimado {datos['nombre']}, tus multas son muy recientes (menos de 3 años). No se pueden borrar.")

    st.markdown("<div style='text-align:center; margin-top:50px; color:#cbd5e1;'>BorraTusMultas.cl - 2026</div>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
